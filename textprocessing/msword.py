from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

from message import Message

def read_docx(fp):
    document = Document(fp)
    return '\n'.join(p.text for p in document.paragraphs)


def write_docx(queue, fp, text, max_center_len=10):
    queue.put(Message('Creating document.'))
    document = Document()
    document.sections[0].left_margin = Inches(.5)
    document.sections[0].right_margin = Inches(.5)

    font = document.styles['Normal'].font
    font.name = 'Lucida Sans Typewriter'
    font.size = Pt(9)   # 9pt	12px	0.75em	75%

    min_center_width = max_center_len * Pt(9)
    side_width = Inches(7.5) - min_center_width

    rows = text.strip().split('\n')

    table = document.add_table(rows=len(rows), cols=3)
    table_cells = table._cells
    table.autofit = True
    col_n = 3

    msg = 'Adding lines to .docx file.'
    queue.put(Message(msg))

    for ri, line in enumerate(rows):
        row_cells = table_cells[ri*col_n:(ri+1)*col_n]
        row = line.split('\t')
        for ci in range(col_n):
            row_cells[ci].alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            row_cells[ci].paragraphs[0].text = row[ci].strip()
            if ci == 0:
                row_cells[ci].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        row_cells[1].width = min_center_width

        row_cells[0].width = side_width
        row_cells[2].width = side_width


    document.save(fp)
    queue.put(Message('Saved concordance as ', fp, tag='red'))